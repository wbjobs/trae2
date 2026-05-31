import os
import re
import io
import time
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from pathlib import Path
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor, as_completed

import pdfplumber
from docx import Document
from loguru import logger
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

from config import settings


@dataclass
class ParseProgress:
    stage: str
    processed: int
    total: int
    last_success_position: int = 0


@dataclass
class ParsedDocument:
    document_id: str
    file_name: str
    file_type: str
    raw_text: str
    cleaned_text: str
    paragraphs: List[str] = field(default_factory=list)
    parties: List[str] = field(default_factory=list)
    case_type: Optional[str] = None
    court: Optional[str] = None
    case_number: Optional[str] = None
    legal_claims: List[str] = field(default_factory=list)
    key_phrases: List[str] = field(default_factory=list)
    metadata: Dict[str, str] = field(default_factory=dict)
    parse_warnings: List[str] = field(default_factory=list)
    is_partial: bool = False


class DocumentParser:
    _instance = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialize()
        return cls._instance

    def _initialize(self):
        self._case_type_patterns = {
            "民事": [r"民事", r"合同", r"侵权", r"婚姻", r"继承", r"劳动", r"借贷", r"买卖"],
            "刑事": [r"刑事", r"犯罪", r"公诉", r"自诉", r"盗窃", r"诈骗", r"故意伤害"],
            "行政": [r"行政", r"行政复议", r"行政诉讼", r"行政处罚"],
            "商事": [r"公司", r"破产", r"票据", r"保险", r"股权", r"合伙"],
            "知识产权": [r"专利", r"商标", r"著作权", r"知识产权", r"侵权"],
        }
        
        self._court_pattern = re.compile(
            r"(最高人民法院|高级人民法院|中级人民法院|基层人民法院|人民法院|海事法院|知识产权法院|互联网法院)"
        )
        
        self._case_number_patterns = [
            re.compile(r"[(（](\d{4})[)）]\s*(\S+)\s*字第\s*(\d+)\s*号"),
            re.compile(r"[(（](\d{4})[)）]\s*(\S+)\s*民初\s*(\d+)\s*号"),
            re.compile(r"[(（](\d{4})[)）]\s*(\S+)\s*刑初\s*(\d+)\s*号"),
            re.compile(r"[(（](\d{4})[)）]\s*(\S+)\s*行初\s*(\d+)\s*号"),
        ]
        
        self._party_patterns = [
            re.compile(r"原告[：:]\s*(\S+?)(?=[，,。；\n被告|被上诉人|第三人])"),
            re.compile(r"被告[：:]\s*(\S+?)(?=[，,。；\n原告|上诉人|第三人])"),
            re.compile(r"上诉人[：:]\s*(\S+?)(?=[，,。；\n被上诉人|原审])"),
            re.compile(r"被上诉人[：:]\s*(\S+?)(?=[，,。；\n上诉人|原审])"),
            re.compile(r"申请人[：:]\s*(\S+?)(?=[，,。；\n被申请人|申请])"),
            re.compile(r"被申请人[：:]\s*(\S+?)(?=[，,。；\n申请人])"),
        ]

        self._chunk_size = 5000
        self._max_chunks = 100
        self._progress: Optional[ParseProgress] = None

        logger.info("DocumentParser initialized with chunked parsing support")

    async def parse_file(
        self, 
        file_content: bytes, 
        file_name: str,
        enable_chunking: bool = True,
        max_retries: int = 3
    ) -> ParsedDocument:
        file_ext = Path(file_name).suffix.lower()
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {file_ext}")

        self._progress = ParseProgress(stage="extracting", processed=0, total=100)
        
        try:
            raw_text = await self._extract_text_with_retry(
                file_content, file_ext, enable_chunking, max_retries
            )
            
            if not raw_text.strip():
                raise ValueError("No text could be extracted from the document")

            self._progress.stage = "cleaning"
            cleaned_text = self._clean_text(raw_text)
            
            self._progress.stage = "splitting"
            paragraphs = self._split_paragraphs_smart(cleaned_text)

            doc_id = self._generate_document_id()

            parsed_doc = ParsedDocument(
                document_id=doc_id,
                file_name=file_name,
                file_type=file_ext,
                raw_text=raw_text[:1000000],
                cleaned_text=cleaned_text[:500000],
                paragraphs=paragraphs,
                is_partial=len(raw_text) > 1000000,
            )

            self._progress.stage = "metadata"
            try:
                self._extract_metadata_safe(parsed_doc)
            except Exception as e:
                parsed_doc.parse_warnings.append(f"Metadata extraction incomplete: {str(e)}")
                logger.warning(f"Metadata extraction failed for {file_name}: {e}")

            self._progress.stage = "claims"
            try:
                self._extract_legal_claims_safe(parsed_doc)
            except Exception as e:
                parsed_doc.parse_warnings.append(f"Claims extraction incomplete: {str(e)}")

            self._progress.stage = "keywords"
            try:
                self._extract_key_phrases_smart(parsed_doc)
            except Exception as e:
                parsed_doc.parse_warnings.append(f"Keyword extraction incomplete: {str(e)}")

            logger.info(
                f"Document parsed successfully: {file_name}, ID: {doc_id}, "
                f"paragraphs: {len(paragraphs)}, warnings: {len(parsed_doc.parse_warnings)}"
            )
            return parsed_doc

        except Exception as e:
            logger.error(f"Failed to parse {file_name}: {e}")
            if 'raw_text' in locals():
                return ParsedDocument(
                    document_id=self._generate_document_id(),
                    file_name=file_name,
                    file_type=file_ext,
                    raw_text=raw_text[:50000] if raw_text else "",
                    cleaned_text=self._clean_text(raw_text[:50000]) if raw_text else "",
                    parse_warnings=[f"Partial parse due to error: {str(e)}"],
                    is_partial=True,
                )
            raise

    async def _extract_text_with_retry(
        self,
        file_content: bytes,
        file_ext: str,
        enable_chunking: bool,
        max_retries: int
    ) -> str:
        @retry(
            stop=stop_after_attempt(max_retries),
            wait=wait_exponential(multiplier=1, min=1, max=10),
            retry=retry_if_exception_type((IOError, OSError, RuntimeError)),
            before_sleep=lambda s: logger.warning(f"Retrying text extraction (attempt {s.attempt_number})")
        )
        async def _extract():
            return await self._extract_text(file_content, file_ext, enable_chunking)
        
        return await _extract()

    async def _extract_text(
        self, 
        file_content: bytes, 
        file_ext: str,
        enable_chunking: bool = True
    ) -> str:
        if file_ext == ".txt":
            return self._safe_decode(file_content)
        elif file_ext in [".docx", ".doc"]:
            return self._extract_from_docx_chunked(file_content, enable_chunking)
        elif file_ext == ".pdf":
            return await self._extract_from_pdf_chunked(file_content, enable_chunking)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    @staticmethod
    def _safe_decode(content: bytes) -> str:
        for encoding in ['utf-8', 'gbk', 'gb2312', 'gb18030', 'utf-16', 'latin-1']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode('utf-8', errors='ignore')

    def _extract_from_docx_chunked(
        self, 
        file_content: bytes, 
        enable_chunking: bool
    ) -> str:
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            total_paragraphs = len(doc.paragraphs)
            
            if enable_chunking and total_paragraphs > 200:
                logger.info(f"Large DOCX detected: {total_paragraphs} paragraphs, using chunked processing")
                
                chunk_size = 100
                for i in range(0, total_paragraphs, chunk_size):
                    chunk = doc.paragraphs[i:i + chunk_size]
                    chunk_text = "\n".join(p.text for p in chunk if p.text.strip())
                    text_parts.append(chunk_text)
                    
                    self._progress.processed = min(i + chunk_size, total_paragraphs)
                    self._progress.total = total_paragraphs
            else:
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

            result = "\n".join(text_parts)
            logger.info(f"DOCX extraction completed: {len(result)} characters")
            return result
            
        except Exception as e:
            logger.error(f"Failed to extract text from docx: {e}")
            return ""

    async def _extract_from_pdf_chunked(
        self, 
        file_content: bytes, 
        enable_chunking: bool
    ) -> str:
        try:
            text_parts = []
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"PDF has {total_pages} pages")
                
                if enable_chunking and total_pages > 20:
                    logger.info(f"Large PDF detected: {total_pages} pages, using chunked processing")
                    
                    batch_size = 10
                    for batch_start in range(0, total_pages, batch_size):
                        batch_end = min(batch_start + batch_size, total_pages)
                        batch_text = []
                        
                        for page_idx in range(batch_start, batch_end):
                            try:
                                page = pdf.pages[page_idx]
                                page_text = page.extract_text() or ""
                                if page_text.strip():
                                    batch_text.append(page_text)
                            except Exception as e:
                                logger.warning(f"Failed to extract page {page_idx}: {e}")
                                continue
                        
                        text_parts.append("\n".join(batch_text))
                        self._progress.processed = batch_end
                        self._progress.total = total_pages
                        
                        await asyncio.sleep(0.01)
                else:
                    for page_idx, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text() or ""
                            if page_text.strip():
                                text_parts.append(page_text)
                        except Exception as e:
                            logger.warning(f"Failed to extract page {page_idx}: {e}")
                            continue

            result = "\n".join(text_parts)
            logger.info(f"PDF extraction completed: {len(result)} characters")
            return result
            
        except Exception as e:
            logger.error(f"Failed to extract text from pdf: {e}")
            return ""

    def _clean_text(self, text: str) -> str:
        if len(text) > 500000:
            logger.warning(f"Text too long ({len(text)} chars), truncating to 500000")
            text = text[:500000]
        
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"[\u3000]+", " ", text)
        text = text.replace("\r\n", "\n")
        text = re.sub(r"\n[ \t]+", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        return text

    def _split_paragraphs_smart(self, text: str) -> List[str]:
        paragraphs = re.split(r"\n\s*\n", text)
        
        merged_paragraphs = []
        current_para = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            if len(para) < 50 and current_para:
                if not para.endswith(("。", "！", "？", "：", ";", ".")):
                    current_para += para
                    continue
            
            if current_para:
                merged_paragraphs.append(current_para)
            current_para = para
        
        if current_para:
            merged_paragraphs.append(current_para)
        
        final_paragraphs = []
        for para in merged_paragraphs:
            if len(para) > 2000:
                sentences = re.split(r"(?<=[。！？.])", para)
                chunk = ""
                for sent in sentences:
                    if len(chunk) + len(sent) < 1500:
                        chunk += sent
                    else:
                        if chunk:
                            final_paragraphs.append(chunk)
                        chunk = sent
                if chunk:
                    final_paragraphs.append(chunk)
            else:
                final_paragraphs.append(para)
        
        return final_paragraphs

    def _extract_metadata_safe(self, doc: ParsedDocument):
        sample_text = doc.cleaned_text[:20000]
        
        court_match = self._court_pattern.search(sample_text)
        if court_match:
            doc.court = court_match.group(0)
        
        for pattern in self._case_number_patterns:
            match = pattern.search(sample_text)
            if match:
                doc.case_number = match.group(0)
                break

        for case_type, patterns in self._case_type_patterns.items():
            score = 0
            for pattern in patterns:
                if re.search(pattern, sample_text):
                    score += 1
            if score >= 2:
                doc.case_type = case_type
                break
            elif score >= 1 and not doc.case_type:
                doc.case_type = case_type

        parties = set()
        for pattern in self._party_patterns:
            try:
                matches = pattern.findall(sample_text)
                for match in matches:
                    if isinstance(match, tuple):
                        match = match[0]
                    party = match.strip("，,。；\n :：")
                    if len(party) > 1 and len(party) < 50 and party not in parties:
                        parties.add(party)
            except Exception as e:
                continue
        doc.parties = list(parties)[:10]

    def _extract_legal_claims_safe(self, doc: ParsedDocument):
        sample_text = doc.cleaned_text[:30000]
        
        claim_patterns = [
            r"诉讼请求[：:](.*?)(?=事实与理由|本院认为|经审理查明|此致)",
            r"请求事项[：:](.*?)(?=事实与理由|本院认为|此致)",
            r"申请事项[：:](.*?)(?=事实与理由|此致)",
            r"仲裁请求[：:](.*?)(?=事实与理由)",
        ]
        
        claims = []
        for pattern in claim_patterns:
            try:
                match = re.search(pattern, sample_text, re.DOTALL)
                if match:
                    claim_text = match.group(1).strip()
                    if claim_text and len(claim_text) > 10:
                        claim_items = re.split(r"(?<=[；;。\n])\s*(?=\d+[、.．:：])", claim_text)
                        for item in claim_items:
                            item = item.strip()
                            if item and len(item) > 5:
                                claims.append(item[:500])
                        break
            except Exception as e:
                continue
        
        doc.legal_claims = claims[:10]

    def _extract_key_phrases_smart(self, doc: ParsedDocument):
        import jieba
        import jieba.analyse

        legal_keywords = [
            "违约责任", "损害赔偿", "合同解除", "合同无效", "缔约过失",
            "不当得利", "无因管理", "侵权责任", "过错责任", "连带责任",
            "诉讼时效", "举证责任", "证据不足", "事实清楚", "证据确实充分",
            "适用法律", "维持原判", "改判", "撤销原判", "发回重审",
            "本金", "利息", "违约金", "赔偿金", "滞纳金",
            "返还财产", "恢复原状", "消除影响", "赔礼道歉",
            "合同成立", "合同生效", "合同终止", "合同撤销",
            "债权", "债务", "担保", "抵押", "质押", "留置", "定金",
            "共同被告", "第三人", "代理人", "辩护人",
            "管辖异议", "回避", "举证期限", "质证", "认证",
        ]

        sample_text = doc.cleaned_text[:20000]
        found_phrases = []
        
        try:
            tfidf_keywords = jieba.analyse.extract_tags(
                sample_text, topK=20, withWeight=False, allowPOS=('n', 'vn', 'v')
            )
            
            for keyword in tfidf_keywords:
                if keyword in legal_keywords:
                    found_phrases.append(keyword)
        except Exception as e:
            logger.debug(f"TF-IDF extraction failed: {e}")

        for keyword in legal_keywords:
            if keyword in sample_text and keyword not in found_phrases:
                found_phrases.append(keyword)
                if len(found_phrases) >= 20:
                    break

        doc.key_phrases = found_phrases[:20]

    @staticmethod
    def _generate_document_id() -> str:
        import uuid
        return f"doc_{uuid.uuid4().hex[:16]}"

    async def parse_batch(
        self, 
        files: List[Tuple[bytes, str]],
        max_workers: int = 4
    ) -> List[ParsedDocument]:
        parsed_docs = []
        
        async def parse_single(content: bytes, name: str) -> Optional[ParsedDocument]:
            try:
                return await self.parse_file(content, name)
            except Exception as e:
                logger.error(f"Failed to parse {name}: {e}")
                return None

        import asyncio
        tasks = [parse_single(content, name) for content, name in files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        for result in results:
            if isinstance(result, ParsedDocument):
                parsed_docs.append(result)
        
        return parsed_docs

    def get_progress(self) -> Optional[ParseProgress]:
        return self._progress


import asyncio
