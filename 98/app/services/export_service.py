import os
import uuid
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

from app.core.config import get_settings
from app.models.task import TaskResult

settings = get_settings()


class ExportService:
    def __init__(self):
        self.export_dir = Path(settings.export_dir)
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def _get_user_export_dir(self, user_id: int) -> Path:
        user_dir = self.export_dir / str(user_id)
        user_dir.mkdir(parents=True, exist_ok=True)
        return user_dir

    def export_to_txt(
        self,
        content: str,
        user_id: int,
        filename: str,
    ) -> str:
        user_dir = self._get_user_export_dir(user_id)
        export_path = user_dir / f"{uuid.uuid4().hex}_{filename}.txt"

        with open(export_path, "w", encoding="utf-8") as f:
            f.write(content)

        return str(export_path)

    def export_to_docx(
        self,
        content: str,
        user_id: int,
        filename: str,
        title: Optional[str] = None,
        corrections: Optional[list] = None,
    ) -> str:
        user_dir = self._get_user_export_dir(user_id)
        export_path = user_dir / f"{uuid.uuid4().hex}_{filename}.docx"

        doc = Document()

        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

        if title:
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            if not para_text.strip():
                continue

            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.3)
            para.paragraph_format.line_spacing = 1.5

            lines = para_text.split("\n")
            for i, line in enumerate(lines):
                if line.strip():
                    run = para.add_run(line.strip())
                    run.font.name = "宋体"
                    run.font.size = Pt(12)
                    if i < len(lines) - 1:
                        run.add_break()

        if corrections:
            doc.add_page_break()
            doc.add_heading("校对修改意见", level=1)

            for i, corr in enumerate(corrections, 1):
                p = doc.add_paragraph()
                p.style = "List Number"

                type_text = {
                    "spelling": "错别字",
                    "grammar": "语法错误",
                    "terminology": "专业术语",
                    "format": "格式问题",
                }.get(corr.get("correction_type"), "其他")

                run = p.add_run(f"【{type_text}】")
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)

                p.add_run(f"\n原文：{corr.get('original_text', '')}\n")
                run = p.add_run(f"建议：{corr.get('corrected_text', '')}\n")
                run.font.color.rgb = RGBColor(0, 128, 0)
                p.add_run(f"说明：{corr.get('explanation', '')}")

        doc.save(export_path)
        return str(export_path)

    def export_to_markdown(
        self,
        content: str,
        user_id: int,
        filename: str,
        title: Optional[str] = None,
        corrections: Optional[list] = None,
    ) -> str:
        user_dir = self._get_user_export_dir(user_id)
        export_path = user_dir / f"{uuid.uuid4().hex}_{filename}.md"

        md_content = []

        if title:
            md_content.append(f"# {title}")
            md_content.append("")

        paragraphs = content.split("\n\n")
        for para in paragraphs:
            if para.strip():
                md_content.append(para.strip())
                md_content.append("")

        if corrections:
            md_content.append("")
            md_content.append("## 校对修改意见")
            md_content.append("")

            for i, corr in enumerate(corrections, 1):
                type_text = {
                    "spelling": "错别字",
                    "grammar": "语法错误",
                    "terminology": "专业术语",
                    "format": "格式问题",
                }.get(corr.get("correction_type"), "其他")

                md_content.append(f"{i}. **【{type_text}】**")
                md_content.append(f"   - 原文：{corr.get('original_text', '')}")
                md_content.append(f"   - 建议：{corr.get('corrected_text', '')}")
                md_content.append(f"   - 说明：{corr.get('explanation', '')}")
                md_content.append("")

        with open(export_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_content))

        return str(export_path)

    def export_to_html(
        self,
        content: str,
        user_id: int,
        filename: str,
        title: Optional[str] = None,
    ) -> str:
        user_dir = self._get_user_export_dir(user_id)
        export_path = user_dir / f"{uuid.uuid4().hex}_{filename}.html"

        html_content = [
            "<!DOCTYPE html>",
            "<html lang='zh-CN'>",
            "<head>",
            "<meta charset='UTF-8'>",
            "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            f"<title>{title or '文档'}</title>",
            "<style>",
            "body { font-family: 'Microsoft YaHei', sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.8; }",
            "h1 { color: #333; text-align: center; border-bottom: 2px solid #007bff; padding-bottom: 10px; }",
            "p { text-indent: 2em; margin-bottom: 1em; }",
            ".correction { background: #fff3cd; padding: 10px; border-left: 4px solid #ffc107; margin: 10px 0; }",
            "</style>",
            "</head>",
            "<body>",
        ]

        if title:
            html_content.append(f"<h1>{title}</h1>")

        md = markdown.markdown(content)
        html_content.append(md)

        html_content.extend(["</body>", "</html>"])

        with open(export_path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_content))

        return str(export_path)

    def export_task_result(
        self,
        result: TaskResult,
        user_id: int,
        format_type: str = "docx",
        filename: Optional[str] = None,
    ) -> str:
        content = result.corrected_content or result.original_content or ""
        base_filename = filename or f"task_{result.task_id}"

        corrections_list = []
        if result.corrections:
            corrections_list = [
                {
                    "correction_type": c.correction_type,
                    "original_text": c.original_text,
                    "corrected_text": c.corrected_text,
                    "explanation": c.explanation,
                }
                for c in result.corrections
            ]

        if format_type == "txt":
            return self.export_to_txt(content, user_id, base_filename)
        elif format_type == "docx":
            return self.export_to_docx(content, user_id, base_filename, base_filename, corrections_list)
        elif format_type == "md":
            return self.export_to_markdown(content, user_id, base_filename, base_filename, corrections_list)
        elif format_type == "html":
            return self.export_to_html(content, user_id, base_filename, base_filename)
        else:
            raise ValueError(f"不支持的导出格式: {format_type}")

    def get_export_file_path(self, file_path: str) -> Optional[str]:
        if os.path.exists(file_path):
            return file_path
        return None


export_service = ExportService()
