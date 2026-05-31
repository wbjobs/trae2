from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from zipfile import BadZipFile
from loguru import logger

from app.parsers.base_parser import BaseParser
from app.schemas.document import DocumentParseResult


class DocxParser(BaseParser):
    def parse(self, file_path: str, password: Optional[str] = None) -> DocumentParseResult:
        try:
            try:
                doc = Document(file_path)
            except PackageNotFoundError:
                return DocumentParseResult(
                    success=False,
                    error="DOCX文件损坏或格式不正确",
                )
            except BadZipFile:
                return DocumentParseResult(
                    success=False,
                    error="文件不是有效的DOCX文档",
                )
            except Exception as e:
                if "password" in str(e).lower() or "encrypted" in str(e).lower():
                    return DocumentParseResult(
                        success=False,
                        error="文档已加密，无法解析。请先解密文档后再上传",
                    )
                raise

            content_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_content.append(cell.text.strip())
                    if row_content:
                        content_parts.append(" | ".join(row_content))

            content = "\n\n".join(content_parts)

            if not content.strip():
                return DocumentParseResult(
                    success=False,
                    error="文档内容为空或无法提取文本",
                )

            return DocumentParseResult(
                success=True,
                content=content,
                word_count=self._count_words(content),
                paragraph_count=self._count_paragraphs(content),
            )

        except Exception as e:
            logger.error(f"Failed to parse DOCX file {file_path}: {e}")
            error_msg = str(e)
            if "password" in error_msg.lower() or "encrypted" in error_msg.lower():
                error_msg = "文档已加密，无法解析。请先解密文档后再上传"
            elif "corrupt" in error_msg.lower() or "damaged" in error_msg.lower():
                error_msg = "文档已损坏，无法解析"
            return DocumentParseResult(
                success=False,
                error=f"DOCX解析失败: {error_msg}",
            )
