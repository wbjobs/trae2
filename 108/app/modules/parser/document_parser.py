import os
import re
import io
import struct
from typing import List, Dict, Any, Optional, Tuple
from abc import ABC, abstractmethod
from zipfile import BadZipFile
import pdfplumber
from pdfminer.pdfparser import PDFSyntaxError
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from app.core import log, BadRequestException


class DocumentParseError(Exception):
    def __init__(self, message: str, error_type: str, recoverable: bool = False, suggestion: str = ""):
        self.message = message
        self.error_type = error_type
        self.recoverable = recoverable
        self.suggestion = suggestion
        super().__init__(self.message)


class EncryptedDocumentError(DocumentParseError):
    def __init__(self, message: str, doc_type: str):
        super().__init__(
            message=message,
            error_type="encrypted",
            recoverable=True,
            suggestion=f"该{doc_type}文档已加密，请提供解密后的文档或使用OCR工具提取文本"
        )


class CorruptedDocumentError(DocumentParseError):
    def __init__(self, message: str, doc_type: str):
        super().__init__(
            message=message,
            error_type="corrupted",
            recoverable=False,
            suggestion=f"{doc_type}文档已损坏，请检查文件完整性"
        )


class UnsupportedFormatError(DocumentParseError):
    def __init__(self, message: str):
        super().__init__(
            message=message,
            error_type="unsupported",
            recoverable=False,
            suggestion="请转换为支持的格式（PDF/Word/文本）"
        )


class EmptyContentError(DocumentParseError):
    def __init__(self, message: str, doc_type: str):
        super().__init__(
            message=message,
            error_type="empty",
            recoverable=True,
            suggestion=f"{doc_type}文档可能是扫描件或图片格式，建议使用OCR工具提取文本"
        )


class BaseParser(ABC):
    def __init__(self):
        self.max_retries = 2
        self.retry_delay = 0.5

    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        pass

    def _check_file_exists(self, file_path: str):
        if not os.path.exists(file_path):
            raise BadRequestException(f"文件不存在: {file_path}")
        if os.path.getsize(file_path) == 0:
            raise EmptyContentError("文件大小为0", "空")

    def _safe_extract(self, extract_func, default=None, retries=2):
        for attempt in range(retries):
            try:
                return extract_func()
            except Exception as e:
                log.warning(f"提取尝试 {attempt + 1} 失败: {str(e)}")
                if attempt == retries - 1:
                    return default
        return default


class PDFParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        log.info(f"解析 PDF 文件: {file_path}")
        self._check_file_exists(file_path)

        text_content = []
        metadata = {}
        warnings = []
        is_encrypted = False

        try:
            encryption_status = self._check_pdf_encryption(file_path)
            if encryption_status["encrypted"]:
                raise EncryptedDocumentError(
                    f"PDF文档已加密（{encryption_status.get('details', '未知加密方式')}）",
                    "PDF"
                )

            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    "num_pages": len(pdf.pages),
                    "title": pdf.metadata.get("Title", ""),
                    "author": pdf.metadata.get("Author", ""),
                    "subject": pdf.metadata.get("Subject", ""),
                    "keywords": pdf.metadata.get("Keywords", ""),
                    "creator": pdf.metadata.get("Creator", ""),
                    "producer": pdf.metadata.get("Producer", ""),
                    "creation_date": str(pdf.metadata.get("CreationDate", "")),
                    "modification_date": str(pdf.metadata.get("ModDate", "")),
                }

                total_pages = len(pdf.pages)
                empty_pages = 0

                for i, page in enumerate(pdf.pages):
                    if page is None:
                        empty_pages += 1
                        continue

                    try:
                        text = self._safe_extract(lambda: page.extract_text())
                        if text and text.strip():
                            text_content.append(text.strip())
                        else:
                            empty_pages += 1
                            if i < 5:
                                warnings.append(f"第{i+1}页可能是图片或扫描件")
                    except Exception as e:
                        log.warning(f"提取第{i+1}页文本失败: {str(e)}")
                        empty_pages += 1
                        continue

                metadata["empty_pages"] = empty_pages
                metadata["text_pages"] = total_pages - empty_pages

                if empty_pages == total_pages and total_pages > 0:
                    raise EmptyContentError(
                        "PDF所有页面均无法提取文本，可能是扫描件或图片格式",
                        "PDF"
                    )
                elif empty_pages > total_pages * 0.5:
                    warnings.append(f"超过50%的页面无法提取文本，可能包含大量扫描页")

        except EncryptedDocumentError:
            raise
        except (PDFSyntaxError, ValueError) as e:
            if "crypt" in str(e).lower() or "encrypt" in str(e).lower():
                raise EncryptedDocumentError("PDF文档已加密或需要密码", "PDF")
            raise CorruptedDocumentError(f"PDF文档损坏: {str(e)}", "PDF")
        except Exception as e:
            error_msg = str(e).lower()
            if "encrypt" in error_msg or "password" in error_msg:
                raise EncryptedDocumentError("PDF文档已加密", "PDF")
            log.error(f"PDF 解析失败: {str(e)}")
            raise CorruptedDocumentError(f"PDF解析失败: {str(e)}", "PDF")

        return {
            "content": "\n\n".join(text_content),
            "pages": text_content,
            "metadata": metadata,
            "warnings": warnings,
            "parse_success": True
        }

    def _check_pdf_encryption(self, file_path: str) -> Dict[str, Any]:
        result = {"encrypted": False, "details": ""}
        try:
            with open(file_path, "rb") as f:
                header = f.read(1024)

                if b"/Encrypt" in header or b"/CryptFilter" in header:
                    result["encrypted"] = True
                    result["details"] = "标准加密"

                if b"/StandardSecurityHandler" in header:
                    result["encrypted"] = True
                    result["details"] = "标准安全处理器加密"

            return result
        except Exception as e:
            log.warning(f"检测PDF加密失败: {str(e)}")
            return result


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        log.info(f"解析 Word 文件: {file_path}")
        self._check_file_exists(file_path)

        paragraphs = []
        tables_content = []
        headers_footers = []
        warnings = []

        try:
            if self._check_doc_encryption(file_path):
                raise EncryptedDocumentError("Word文档已加密或受保护", "Word")

            try:
                doc = DocxDocument(file_path)
            except PackageNotFoundError:
                raise CorruptedDocumentError("Word文档损坏或不是有效的.docx文件", "Word")
            except BadZipFile:
                raise CorruptedDocumentError("Word文档不是有效的ZIP压缩包格式", "Word")
            except Exception as e:
                error_msg = str(e).lower()
                if "encrypted" in error_msg or "password" in error_msg or "protect" in error_msg:
                    raise EncryptedDocumentError("Word文档已加密或受保护", "Word")
                raise

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        "text": text,
                        "style": para.style.name,
                        "level": self._get_heading_level(para.style.name),
                        "alignment": str(para.alignment) if para.alignment else "left"
                    })

            for table_idx, table in enumerate(doc.tables):
                try:
                    table_data = []
                    for row_idx, row in enumerate(table.rows):
                        row_data = []
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            row_data.append(cell_text)
                        if any(row_data):
                            table_data.append(row_data)
                    if table_data:
                        tables_content.append({
                            "index": table_idx,
                            "data": table_data,
                            "rows": len(table_data),
                            "cols": max(len(row) for row in table_data) if table_data else 0
                        })
                except Exception as e:
                    log.warning(f"提取表格{table_idx}失败: {str(e)}")
                    continue

            try:
                for section in doc.sections:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            headers_footers.append({"type": "header", "text": paragraph.text.strip()})
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            headers_footers.append({"type": "footer", "text": paragraph.text.strip()})
            except Exception as e:
                log.warning(f"提取页眉页脚失败: {str(e)}")

            content = "\n".join([p["text"] for p in paragraphs])
            if not content.strip() and not tables_content:
                raise EmptyContentError("Word文档中未提取到任何文本内容", "Word")

        except EncryptedDocumentError:
            raise
        except CorruptedDocumentError:
            raise
        except Exception as e:
            log.error(f"Word 解析失败: {str(e)}")
            raise CorruptedDocumentError(f"Word解析失败: {str(e)}", "Word")

        return {
            "content": content,
            "paragraphs": paragraphs,
            "tables": tables_content,
            "headers_footers": headers_footers,
            "warnings": warnings,
            "metadata": {
                "num_paragraphs": len(paragraphs),
                "num_tables": len(tables_content),
                "num_headers_footers": len(headers_footers)
            },
            "parse_success": True
        }

    def _check_doc_encryption(self, file_path: str) -> bool:
        try:
            with open(file_path, "rb") as f:
                header = f.read(512)

                if header.startswith(b'\xd0\xcf\x11\xe0'):
                    return True

                if b"Encryption" in header or b"DataSpaces" in header:
                    return True

            file_size = os.path.getsize(file_path)
            if file_size > 100 and file_size < 1024:
                with open(file_path, "rb") as f:
                    content = f.read()
                    if b"[Content_Types].xml" not in content:
                        return True

            return False
        except Exception as e:
            log.warning(f"检测Word加密失败: {str(e)}")
            return False

    def _get_heading_level(self, style_name: str) -> Optional[int]:
        match = re.match(r"Heading (\d+)", style_name, re.IGNORECASE)
        if match:
            return int(match.group(1))
        match = re.match(r"标题\s*(\d+)", style_name)
        if match:
            return int(match.group(1))
        return None


class TextParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        log.info(f"解析文本文件: {file_path}")
        self._check_file_exists(file_path)

        encodings = [
            "utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030",
            "utf-16", "utf-16-le", "utf-16-be",
            "big5", "shift_jis", "euc-jp", "cp936", "latin1"
        ]

        content = None
        used_encoding = None
        encoding_scores = {}

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    test_content = f.read(4096)
                    score = self._score_encoding(test_content, encoding)
                    encoding_scores[encoding] = score
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                log.debug(f"编码 {encoding} 测试失败: {str(e)}")
                continue

        if encoding_scores:
            best_encoding = max(encoding_scores.keys(), key=lambda x: encoding_scores[x])
            try:
                with open(file_path, "r", encoding=best_encoding) as f:
                    content = f.read()
                    used_encoding = best_encoding
            except Exception as e:
                log.warning(f"使用最佳编码 {best_encoding} 读取失败: {str(e)}")

        if content is None:
            log.warning("所有标准编码均失败，尝试使用容错模式读取")
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()
                    used_encoding = "utf-8 (replacement mode)"
            except Exception as e:
                raise UnsupportedFormatError(f"无法识别文件编码: {str(e)}")

        lines = [line.strip() for line in content.split("\n") if line.strip()]

        if not content.strip():
            raise EmptyContentError("文本文件为空或仅包含空白字符", "文本")

        return {
            "content": content,
            "lines": lines,
            "encoding_scores": encoding_scores,
            "metadata": {
                "encoding": used_encoding,
                "num_lines": len(lines),
                "num_chars": len(content),
                "file_size": os.path.getsize(file_path)
            },
            "parse_success": True
        }

    def _score_encoding(self, content: str, encoding: str) -> float:
        if not content:
            return 0.0

        score = 1.0

        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', content))
        if chinese_chars > 0:
            chinese_ratio = chinese_chars / len(content)
            if encoding in ["gbk", "gb2312", "gb18030", "cp936"]:
                score += chinese_ratio * 2
            elif encoding in ["utf-8", "utf-8-sig"]:
                score += chinese_ratio * 1.5

        null_chars = content.count('\x00')
        if null_chars > 0:
            score -= null_chars * 0.1

        control_chars = len(re.findall(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', content))
        if control_chars > 0:
            score -= control_chars * 0.05

        return max(0.0, score)


class ParserFactory:
    _parsers = {
        ".pdf": PDFParser,
        ".docx": DocxParser,
        ".doc": DocxParser,
        ".txt": TextParser,
        ".text": TextParser,
        ".md": TextParser,
    }

    @classmethod
    def get_parser(cls, file_extension: str) -> BaseParser:
        file_extension = file_extension.lower()
        parser_class = cls._parsers.get(file_extension)
        if not parser_class:
            raise UnsupportedFormatError(f"不支持的文件类型: {file_extension}")
        return parser_class()

    @classmethod
    def supported_extensions(cls) -> List[str]:
        return list(cls._parsers.keys())


def parse_document(file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
    if not os.path.exists(file_path):
        raise BadRequestException(f"文件不存在: {file_path}")

    if not file_type:
        _, file_ext = os.path.splitext(file_path)
    else:
        file_ext = file_type

    parser = ParserFactory.get_parser(file_ext)

    try:
        result = parser.parse(file_path)
        log.info(f"文档解析成功: {file_path}, 内容长度: {len(result.get('content', ''))}")
        return result
    except DocumentParseError as e:
        log.warning(f"文档解析异常 [{e.error_type}]: {e.message}")
        return {
            "content": "",
            "pages": [],
            "metadata": {},
            "error": {
                "type": e.error_type,
                "message": e.message,
                "recoverable": e.recoverable,
                "suggestion": e.suggestion
            },
            "parse_success": False
        }
    except Exception as e:
        log.error(f"文档解析失败: {str(e)}")
        return {
            "content": "",
            "pages": [],
            "metadata": {},
            "error": {
                "type": "unknown",
                "message": str(e),
                "recoverable": False,
                "suggestion": "请检查文件格式是否正确"
            },
            "parse_success": False
        }
