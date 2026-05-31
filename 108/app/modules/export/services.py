import os
import json
from typing import List, Dict, Any, Optional
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from app.core import settings, log, NotFoundException
from app.models import Task, ComparisonResult, Law, Case


class ExportService:
    @staticmethod
    async def export_comparison_results(
        db,
        task_id: int,
        export_format: str = "excel",
        include_analysis: bool = True
    ) -> str:
        task = await db.get(Task, task_id)
        if not task:
            raise NotFoundException("任务不存在")

        results, total = await ExportService._get_task_results(db, task_id)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"比对结果_{task.name}_{timestamp}.{export_format}"
        filepath = os.path.join(settings.EXPORT_DIR, filename)
        os.makedirs(settings.EXPORT_DIR, exist_ok=True)

        if export_format == "excel":
            await ExportService._export_to_excel(results, filepath, task, include_analysis)
        elif export_format == "pdf":
            await ExportService._export_to_pdf(results, filepath, task, include_analysis)
        elif export_format == "json":
            await ExportService._export_to_json(results, filepath, task)
        else:
            raise ValueError(f"不支持的导出格式: {export_format}")

        log.info(f"导出完成: {filepath}")
        return filename

    @staticmethod
    async def export_search_results(
        results: List[Dict[str, Any]],
        search_type: str,
        keyword: str,
        export_format: str = "excel"
    ) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"检索结果_{search_type}_{timestamp}.{export_format}"
        filepath = os.path.join(settings.EXPORT_DIR, filename)
        os.makedirs(settings.EXPORT_DIR, exist_ok=True)

        task_info = {
            "name": f"{search_type}检索",
            "created_at": datetime.now().isoformat(),
            "params": {"keyword": keyword, "search_type": search_type}
        }

        if export_format == "excel":
            await ExportService._export_search_to_excel(results, filepath, task_info, search_type)
        elif export_format == "pdf":
            await ExportService._export_search_to_pdf(results, filepath, task_info, search_type)
        elif export_format == "json":
            await ExportService._export_to_json(results, filepath, task_info)
        else:
            raise ValueError(f"不支持的导出格式: {export_format}")

        log.info(f"检索结果导出完成: {filepath}")
        return filename

    @staticmethod
    async def generate_comparison_report(
        case_content: str,
        matched_laws: List[Dict[str, Any]],
        export_format: str = "pdf",
        db=None
    ) -> str:
        from app.modules.ai import AIService

        report_data = await AIService.generate_comparison_report(case_content, matched_laws)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"比对分析报告_{timestamp}.{export_format}"
        filepath = os.path.join(settings.EXPORT_DIR, filename)
        os.makedirs(settings.EXPORT_DIR, exist_ok=True)

        if export_format == "pdf":
            await ExportService._export_report_to_pdf(report_data, filepath, case_content, matched_laws)
        elif export_format == "docx":
            await ExportService._export_report_to_docx(report_data, filepath, case_content, matched_laws)
        else:
            raise ValueError(f"不支持的报告格式: {export_format}")

        log.info(f"分析报告生成完成: {filepath}")
        return filename

    @staticmethod
    async def _get_task_results(db, task_id: int):
        from app.modules.tasks import TaskService
        return await TaskService.get_task_results(db, task_id, limit=10000)

    @staticmethod
    async def _export_to_excel(
        results: List[Dict[str, Any]],
        filepath: str,
        task: Task,
        include_analysis: bool
    ):
        wb = Workbook()
        ws = wb.active
        ws.title = "比对结果"

        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left_align = Alignment(horizontal="left", vertical="top", wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        headers = ["序号", "案例标题", "法条标题", "法条编号", "相似度(%)", "风险等级"]
        if include_analysis:
            headers.extend(["匹配分析", "关键点", "法律建议"])

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border

        for row_idx, result in enumerate(results, 2):
            score = result.get("similarity_score", 0)
            risk_level = ExportService._get_risk_level(score)

            data = [
                row_idx - 1,
                result.get("case_title", ""),
                result.get("law_title", ""),
                result.get("law_article_no", ""),
                score,
                risk_level
            ]
            if include_analysis:
                data.extend([
                    result.get("matching_analysis", ""),
                    ", ".join(result.get("key_points", [])),
                    result.get("recommendations", "")
                ])

            for col, value in enumerate(data, 1):
                cell = ws.cell(row=row_idx, column=col, value=value)
                cell.border = thin_border
                cell.alignment = left_align

                if col == 5 and isinstance(value, (int, float)):
                    if value >= 80:
                        cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                    elif value >= 60:
                        cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")

        column_widths = [8, 30, 35, 15, 12, 12]
        if include_analysis:
            column_widths.extend([50, 30, 30])
        for col, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = width

        ws.freeze_panes = "A2"

        ws2 = wb.create_sheet("任务信息")
        info_data = [
            ["任务名称", task.name],
            ["任务类型", task.task_type],
            ["创建时间", task.created_at.strftime("%Y-%m-%d %H:%M:%S") if task.created_at else ""],
            ["开始时间", task.started_at.strftime("%Y-%m-%d %H:%M:%S") if task.started_at else ""],
            ["完成时间", task.completed_at.strftime("%Y-%m-%d %H:%M:%S") if task.completed_at else ""],
            ["状态", task.status],
            ["总数", task.total],
            ["成功数", task.completed],
            ["失败数", task.failed],
            ["进度", f"{task.progress}%"]
        ]
        for row, (key, value) in enumerate(info_data, 1):
            ws2.cell(row=row, column=1, value=key).font = Font(bold=True)
            ws2.cell(row=row, column=2, value=str(value))

        wb.save(filepath)

    @staticmethod
    async def _export_to_pdf(
        results: List[Dict[str, Any]],
        filepath: str,
        task: Task,
        include_analysis: bool
    ):
        ExportService._register_fonts()

        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                rightMargin=2 * cm, leftMargin=2 * cm,
                                topMargin=2 * cm, bottomMargin=2 * cm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontName='SimSun', fontSize=18, spaceAfter=12)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontName='SimSun', fontSize=14, spaceAfter=10)
        normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontName='SimSun', fontSize=10, leading=14)
        small_style = ParagraphStyle('CustomSmall', parent=styles['Normal'], fontName='SimSun', fontSize=9, leading=12)

        story = []

        story.append(Paragraph("法律案例比对结果报告", title_style))
        story.append(Spacer(1, 0.5 * cm))

        info_data = [
            ["任务名称:", task.name],
            ["任务类型:", task.task_type],
            ["创建时间:", task.created_at.strftime("%Y-%m-%d %H:%M:%S") if task.created_at else ""],
            ["状态:", task.status],
            ["总记录数:", str(len(results))],
        ]

        info_table = Table(info_data, colWidths=[3 * cm, 10 * cm])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('FONTNAME', (0, 0), (0, -1), 'SimSun'),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.8 * cm))

        for idx, result in enumerate(results, 1):
            story.append(Paragraph(f"{idx}. {result.get('case_title', '')}", heading_style))

            score = result.get("similarity_score", 0)
            risk_level = ExportService._get_risk_level(score)

            basic_info = [
                ["匹配法条:", result.get("law_title", "")],
                ["法条编号:", result.get("law_article_no", "")],
                ["相似度:", f"{score}%"],
                ["风险等级:", risk_level],
            ]
            basic_table = Table(basic_info, colWidths=[3 * cm, 13 * cm])
            basic_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (1, 2), (1, 2), colors.lightcoral if score >= 80 else (colors.lightyellow if score >= 60 else colors.lightgreen)),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(basic_table)

            if include_analysis:
                story.append(Spacer(1, 0.3 * cm))
                story.append(Paragraph("<b>匹配分析:</b>", small_style))
                story.append(Paragraph(result.get("matching_analysis", "")[:1000], normal_style))

                if result.get("key_points"):
                    story.append(Spacer(1, 0.2 * cm))
                    story.append(Paragraph("<b>关键点:</b>", small_style))
                    key_points_text = "、".join(result.get("key_points", []))
                    story.append(Paragraph(key_points_text, normal_style))

                story.append(Spacer(1, 0.2 * cm))
                story.append(Paragraph("<b>法律建议:</b>", small_style))
                story.append(Paragraph(result.get("recommendations", "")[:500], normal_style))

            story.append(Spacer(1, 0.5 * cm))

            if idx % 3 == 0 and idx < len(results):
                story.append(PageBreak())

        doc.build(story)

    @staticmethod
    async def _export_to_json(results: List[Dict[str, Any]], filepath: str, task_info):
        export_data = {
            "export_time": datetime.now().isoformat(),
            "task_info": {
                "name": getattr(task_info, "name", ""),
                "task_type": getattr(task_info, "task_type", ""),
                "created_at": getattr(task_info, "created_at", "").isoformat() if getattr(task_info, "created_at", None) else None,
                "params": getattr(task_info, "params", {})
            },
            "total_results": len(results),
            "results": results
        }

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)

    @staticmethod
    async def _export_search_to_excel(
        results: List[Dict[str, Any]],
        filepath: str,
        task_info: Dict[str, Any],
        search_type: str
    ):
        wb = Workbook()
        ws = wb.active
        ws.title = "检索结果"

        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left_align = Alignment(horizontal="left", vertical="top", wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        if search_type == "law":
            headers = ["序号", "标题", "法条编号", "类型", "类别", "章节", "相关度", "内容摘要"]
        else:
            headers = ["序号", "标题", "案号", "法院", "案件类型", "相关度", "内容摘要"]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border

        for row_idx, result in enumerate(results, 2):
            if search_type == "law":
                data = [
                    row_idx - 1,
                    result.get("title", ""),
                    result.get("article_no", ""),
                    result.get("law_type", ""),
                    result.get("category", ""),
                    result.get("chapter", ""),
                    f"{result.get('_score', 0):.2f}",
                    result.get("content", "")[:200]
                ]
            else:
                data = [
                    row_idx - 1,
                    result.get("title", ""),
                    result.get("case_no", ""),
                    result.get("court", ""),
                    result.get("case_type", ""),
                    f"{result.get('_score', 0):.2f}",
                    (result.get("summary", "") or result.get("content", ""))[:200]
                ]

            for col, value in enumerate(data, 1):
                cell = ws.cell(row=row_idx, column=col, value=value)
                cell.border = thin_border
                cell.alignment = left_align

        column_widths = [8, 30, 15, 12, 12, 15, 12, 50]
        for col, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = width

        ws.freeze_panes = "A2"
        wb.save(filepath)

    @staticmethod
    async def _export_search_to_pdf(
        results: List[Dict[str, Any]],
        filepath: str,
        task_info: Dict[str, Any],
        search_type: str
    ):
        ExportService._register_fonts()

        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                rightMargin=2 * cm, leftMargin=2 * cm,
                                topMargin=2 * cm, bottomMargin=2 * cm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontName='SimSun', fontSize=18, spaceAfter=12)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontName='SimSun', fontSize=14, spaceAfter=10)
        normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontName='SimSun', fontSize=10, leading=14)

        story = []
        story.append(Paragraph(f"{'法条' if search_type == 'law' else '案例'}检索结果报告", title_style))
        story.append(Spacer(1, 0.5 * cm))

        for idx, result in enumerate(results, 1):
            title = result.get("title", "")
            score = result.get("_score", 0)
            story.append(Paragraph(f"{idx}. {title} (相关度: {score:.2f})", heading_style))

            if search_type == "law":
                info = f"编号: {result.get('article_no', '')} | 类型: {result.get('law_type', '')} | 类别: {result.get('category', '')}"
            else:
                info = f"案号: {result.get('case_no', '')} | 法院: {result.get('court', '')} | 类型: {result.get('case_type', '')}"
            story.append(Paragraph(info, normal_style))
            story.append(Spacer(1, 0.2 * cm))

            content = result.get("summary", "") or result.get("content", "")
            story.append(Paragraph(content[:500], normal_style))
            story.append(Spacer(1, 0.5 * cm))

            if idx % 5 == 0 and idx < len(results):
                story.append(PageBreak())

        doc.build(story)

    @staticmethod
    async def _export_report_to_pdf(
        report_data: Dict[str, Any],
        filepath: str,
        case_content: str,
        matched_laws: List[Dict[str, Any]]
    ):
        ExportService._register_fonts()

        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                rightMargin=2 * cm, leftMargin=2 * cm,
                                topMargin=2 * cm, bottomMargin=2 * cm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontName='SimSun', fontSize=18, spaceAfter=12)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontName='SimSun', fontSize=14, spaceAfter=10)
        normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontName='SimSun', fontSize=10, leading=16)

        story = []
        story.append(Paragraph("法律案例比对分析报告", title_style))
        story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
        story.append(Paragraph(f"风险等级: {report_data.get('risk_level', '未知')}", normal_style))
        story.append(Spacer(1, 0.5 * cm))

        story.append(Paragraph("一、案例摘要", heading_style))
        story.append(Paragraph(case_content[:1000].replace('\n', '<br/>'), normal_style))
        story.append(Spacer(1, 0.5 * cm))

        story.append(Paragraph("二、相关法条分析", heading_style))
        for idx, law in enumerate(matched_laws[:5], 1):
            story.append(Paragraph(f"{idx}. {law.get('title', '')} (相似度: {law.get('similarity_score', 0)}%)",
                                   ParagraphStyle('SubHeading', parent=heading_style, fontSize=12)))
            story.append(Paragraph(f"法条内容: {law.get('content', '')[:300]}", normal_style))
            if law.get("matching_analysis"):
                story.append(Paragraph(f"匹配分析: {law.get('matching_analysis', '')[:500]}", normal_style))
            story.append(Spacer(1, 0.3 * cm))

        story.append(PageBreak())
        story.append(Paragraph("三、AI 分析报告", heading_style))
        report_text = report_data.get("report", "").replace('\n', '<br/>')
        for paragraph in report_text.split('<br/>'):
            if paragraph.strip():
                story.append(Paragraph(paragraph, normal_style))
                story.append(Spacer(1, 0.1 * cm))

        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph("四、处理建议", heading_style))
        story.append(Paragraph(report_data.get("summary", "").replace('\n', '<br/>'), normal_style))

        doc.build(story)

    @staticmethod
    async def _export_report_to_docx(
        report_data: Dict[str, Any],
        filepath: str,
        case_content: str,
        matched_laws: List[Dict[str, Any]]
    ):
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading("法律案例比对分析报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"风险等级: {report_data.get('risk_level', '未知')}")

        doc.add_heading("一、案例摘要", level=1)
        doc.add_paragraph(case_content[:1000])

        doc.add_heading("二、相关法条分析", level=1)
        for idx, law in enumerate(matched_laws[:5], 1):
            doc.add_heading(f"{idx}. {law.get('title', '')} (相似度: {law.get('similarity_score', 0)}%)", level=2)
            doc.add_paragraph(f"法条内容: {law.get('content', '')[:300]}")
            if law.get("matching_analysis"):
                doc.add_paragraph(f"匹配分析: {law.get('matching_analysis', '')[:500]}")

        doc.add_heading("三、AI 分析报告", level=1)
        doc.add_paragraph(report_data.get("report", ""))

        doc.add_heading("四、处理建议", level=1)
        doc.add_paragraph(report_data.get("summary", ""))

        doc.save(filepath)

    @staticmethod
    def _register_fonts():
        font_paths = [
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/System/Library/Fonts/PingFang.ttc"
        ]

        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('SimSun', font_path))
                    return
                except Exception:
                    continue

    @staticmethod
    def _get_risk_level(score: int) -> str:
        if score >= 80:
            return "高风险"
        elif score >= 60:
            return "中风险"
        elif score >= 40:
            return "低风险"
        else:
            return "无风险"
