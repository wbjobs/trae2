import os
import subprocess
import tempfile
from loguru import logger
from app.document_parser.base import BaseParser, ParseError, UnsupportedFormatError


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".doc":
            return self._parse_doc_format(file_path)
        return self._parse_docx(file_path)

    def _parse_docx(self, file_path: str) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ParseError("python-docx not installed")

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ParseError(f"Cannot open DOCX file: {e}")

        content_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = (para.style.name or "").lower() if para.style else ""
                if "heading" in style_name:
                    level = 1
                    try:
                        level = int("".join(c for c in style_name if c.isdigit()) or "1")
                    except ValueError:
                        pass
                    content_parts.append(f"{'#' * level} {text}")
                else:
                    content_parts.append(text)

        if doc.tables:
            content_parts.append("--- 表格内容 ---")
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_text = self._extract_table(table)
                    if table_text:
                        content_parts.append(f"表格{table_idx + 1}:")
                        content_parts.append(table_text)
                except Exception as e:
                    logger.warning(f"Failed to extract table {table_idx + 1}: {e}")

        try:
            for section in doc.sections:
                header = section.header
                if header and header.paragraphs:
                    header_text = "\n".join(
                        p.text.strip() for p in header.paragraphs if p.text.strip()
                    )
                    if header_text:
                        content_parts.append(f"[页眉] {header_text}")

                footer = section.footer
                if footer and footer.paragraphs:
                    footer_text = "\n".join(
                        p.text.strip() for p in footer.paragraphs if p.text.strip()
                    )
                    if footer_text:
                        content_parts.append(f"[页脚] {footer_text}")
        except Exception as e:
            logger.debug(f"Header/footer extraction skipped: {e}")

        content = "\n\n".join(content_parts)
        if not content.strip():
            raise ParseError(f"No text content extracted from DOCX: {file_path}")
        return content

    def _extract_table(self, table) -> str:
        rows_text = []
        seen_rows = set()

        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                cell_content = cell.text.strip()
                paragraphs = []
                for p in cell.paragraphs:
                    pt = p.text.strip()
                    if pt:
                        paragraphs.append(pt)
                if paragraphs:
                    cell_content = "\n".join(paragraphs)
                cells_text.append(cell_content)

            row_text = " | ".join(cells_text)
            row_key = row_text.strip("| ")
            if row_key and row_key not in seen_rows:
                seen_rows.add(row_key)
                rows_text.append(row_text)

        return "\n".join(rows_text)

    def _parse_doc_format(self, file_path: str) -> str:
        try:
            result = self._convert_doc_to_docx(file_path)
            if result:
                return result
        except Exception as e:
            logger.warning(f"DOC conversion failed: {e}")

        try:
            text = self._extract_doc_with_antiword(file_path)
            if text:
                return text
        except Exception as e:
            logger.debug(f"antiword extraction failed: {e}")

        raise ParseError(
            f"Cannot parse .doc file (old Word format): {file_path}. "
            f"Please convert to .docx format first."
        )

    def _convert_doc_to_docx(self, file_path: str) -> str | None:
        try:
            import subprocess
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "docx",
                     "--outdir", os.path.dirname(tmp_path), file_path],
                    capture_output=True, timeout=60,
                )
                converted_name = os.path.splitext(os.path.basename(file_path))[0] + ".docx"
                converted_path = os.path.join(os.path.dirname(tmp_path), converted_name)
                if os.path.exists(converted_path):
                    result = self._parse_docx(converted_path)
                    os.remove(converted_path)
                    return result
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        return None

    def _extract_doc_with_antiword(self, file_path: str) -> str | None:
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        return None
