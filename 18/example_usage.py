#!/usr/bin/env python3
"""
文档语义抽取与智能归类AI系统 - 使用示例
"""

import os
import sys
import asyncio
import httpx

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

BASE_URL = "http://localhost:8000"
API_PREFIX = "/api/v1"


async def example_upload_document():
    """示例：上传文档"""
    print("=" * 50)
    print("示例1: 上传文档")
    print("=" * 50)

    from docx import Document
    import tempfile

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('技术服务合同', level=1)
        doc.add_paragraph('甲方：XX科技有限公司')
        doc.add_paragraph('乙方：YY软件有限公司')
        doc.add_paragraph('根据《中华人民共和国合同法》，甲乙双方就技术服务事宜达成协议。')
        doc.add_paragraph('服务内容：企业管理系统开发与实施')
        doc.add_paragraph('服务期限：2024年1月1日至2024年12月31日')
        doc.add_paragraph('合同金额：人民币50万元整')
        doc.save(tmp.name)
        tmp_path = tmp.name

    try:
        async with httpx.AsyncClient() as client:
            with open(tmp_path, 'rb') as f:
                files = {'file': ('test_contract.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                response = await client.post(f"{BASE_URL}{API_PREFIX}/documents/upload", files=files)

            if response.status_code == 201:
                doc_info = response.json()
                print(f"上传成功!")
                print(f"文档ID: {doc_info['id']}")
                print(f"文件名: {doc_info['filename']}")
                print(f"文件大小: {doc_info['file_size']} 字节")
                print(f"状态: {doc_info['status']}")
                return doc_info['id']
            else:
                print(f"上传失败: {response.status_code} - {response.text}")
                return None
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def example_process_document(document_id: int):
    """示例：处理文档"""
    print("\n" + "=" * 50)
    print("示例2: 处理文档")
    print("=" * 50)

    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{BASE_URL}{API_PREFIX}/documents/{document_id}/process"
        )

        if response.status_code == 200:
            doc_info = response.json()
            print(f"处理任务已启动!")
            print(f"文档ID: {doc_info['id']}")
            print(f"当前状态: {doc_info['status']}")
            print("文档正在后台处理中，请稍后获取结果...")
        else:
            print(f"启动处理失败: {response.status_code} - {response.text}")


async def example_get_document(document_id: int):
    """示例：获取文档详情"""
    print("\n" + "=" * 50)
    print("示例3: 获取文档详情")
    print("=" * 50)

    async with httpx.AsyncClient() as client:
        response = await client.get(f"{BASE_URL}{API_PREFIX}/documents/{document_id}")

        if response.status_code == 200:
            doc = response.json()
            print(f"文档ID: {doc['id']}")
            print(f"文件名: {doc['document_info']['filename']}")
            print(f"状态: {doc['document_info']['status']}")
            print(f"上传时间: {doc['document_info']['upload_time']}")

            if doc.get('semantic_features'):
                print(f"\n语义特征:")
                print(f"  关键词: {doc['semantic_features']['keywords'][:10]}")
                print(f"  摘要: {doc['semantic_features']['summary'][:100]}...")
                print(f"  主题: {doc['semantic_features']['topics']}")
                print(f"  情感: {doc['semantic_features']['sentiment']}")

            if doc.get('classification'):
                print(f"\n分类结果:")
                print(f"  主分类: {doc['classification']['primary_category']}")
                print(f"  次要分类: {doc['classification']['secondary_categories']}")
                print(f"  置信度: {doc['classification']['confidence']:.2f}")
        else:
            print(f"获取文档失败: {response.status_code} - {response.text}")


async def example_list_documents():
    """示例：获取文档列表"""
    print("\n" + "=" * 50)
    print("示例4: 获取文档列表")
    print("=" * 50)

    async with httpx.AsyncClient() as client:
        response = await client.get(f"{BASE_URL}{API_PREFIX}/documents", params={"limit": 10})

        if response.status_code == 200:
            docs = response.json()
            print(f"共找到 {len(docs)} 个文档:")
            for doc in docs[:5]:
                category = doc.get('classification', {}).get('primary_category', '未分类') if doc.get('classification') else '未分类'
                print(f"  - ID:{doc['id']} {doc['document_info']['filename']} [{doc['document_info']['status']}] - {category}")
        else:
            print(f"获取列表失败: {response.status_code} - {response.text}")


async def example_search_documents():
    """示例：语义搜索文档"""
    print("\n" + "=" * 50)
    print("示例5: 语义搜索文档")
    print("=" * 50)

    search_query = {
        "query": "合同协议 技术服务",
        "top_k": 5,
        "categories": ["合同协议", "技术文档"]
    }

    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{BASE_URL}{API_PREFIX}/documents/search",
            json=search_query
        )

        if response.status_code == 200:
            results = response.json()
            print(f"搜索关键词: {search_query['query']}")
            print(f"找到 {len(results)} 个相关文档:")
            for result in results:
                print(f"  - [{result['similarity_score']:.3f}] {result['filename']} ({result['category']})")
                if result.get('matched_keywords'):
                    print(f"    匹配关键词: {result['matched_keywords']}")
        else:
            print(f"搜索失败: {response.status_code} - {response.text}")


async def example_batch_process():
    """示例：批量处理文档"""
    print("\n" + "=" * 50)
    print("示例6: 批量处理文档")
    print("=" * 50)

    batch_request = {
        "document_ids": [1, 2, 3],
        "skip_classification": False,
        "skip_embedding": False
    }

    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{BASE_URL}{API_PREFIX}/documents/batch/process",
            json=batch_request
        )

        if response.status_code == 200:
            result = response.json()
            print(f"批量任务已创建!")
            print(f"任务ID: {result['task_id']}")
            print(f"文档总数: {result['total_count']}")
            print(f"状态: {result['status']}")

            await asyncio.sleep(2)

            status_response = await client.get(
                f"{BASE_URL}{API_PREFIX}/documents/batch/{result['task_id']}/status"
            )
            if status_response.status_code == 200:
                status = status_response.json()
                print(f"\n任务状态:")
                print(f"  状态: {status['status']}")
                print(f"  已处理: {status['processed_count']}/{status['total_count']}")
                print(f"  失败: {status['failed_count']}")
        else:
            print(f"批量处理失败: {response.status_code} - {response.text}")


async def example_call_external_system():
    """示例：调用外部系统"""
    print("\n" + "=" * 50)
    print("示例7: 调用外部系统接口")
    print("=" * 50)

    external_call = {
        "system_name": "示例API",
        "endpoint": "https://jsonplaceholder.typicode.com/posts/1",
        "method": "GET",
        "timeout": 30
    }

    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{BASE_URL}{API_PREFIX}/external/call",
            json=external_call
        )

        if response.status_code == 200:
            result = response.json()
            print(f"调用成功: {result['success']}")
            print(f"状态码: {result['status_code']}")
            print(f"响应时间: {result['response_time']:.3f}s")
            if result.get('response_data'):
                print(f"响应数据: {str(result['response_data'])[:100]}...")
        else:
            print(f"调用失败: {response.status_code} - {response.text}")


async def example_get_categories():
    """示例：获取分类列表"""
    print("\n" + "=" * 50)
    print("示例8: 获取支持的分类列表")
    print("=" * 50)

    async with httpx.AsyncClient() as client:
        response = await client.get(f"{BASE_URL}{API_PREFIX}/categories")

        if response.status_code == 200:
            data = response.json()
            print(f"支持的分类 ({len(data['categories'])} 个):")
            for cat in data['categories']:
                print(f"  - {cat}")
        else:
            print(f"获取分类失败: {response.status_code} - {response.text}")


async def example_health_check():
    """示例：健康检查"""
    print("\n" + "=" * 50)
    print("示例9: 健康检查")
    print("=" * 50)

    async with httpx.AsyncClient() as client:
        response = await client.get(f"{BASE_URL}{API_PREFIX}/health")

        if response.status_code == 200:
            data = response.json()
            print(f"服务状态: {data['status']}")
            print(f"服务名称: {data['service']}")
            print(f"版本: {data['version']}")
        else:
            print(f"健康检查失败: {response.status_code} - {response.text}")


async def main():
    """运行所有示例"""
    print("文档语义抽取与智能归类AI系统 - API使用示例")
    print(f"服务地址: {BASE_URL}")
    print("请确保服务已启动: python main.py")
    print()

    try:
        await example_health_check()
        await example_get_categories()

        doc_id = await example_upload_document()
        if doc_id:
            await example_process_document(doc_id)

            print("\n等待处理完成...")
            await asyncio.sleep(3)

            await example_get_document(doc_id)

        await example_list_documents()
        await example_search_documents()
        await example_call_external_system()

        print("\n" + "=" * 50)
        print("所有示例执行完成!")
        print("=" * 50)

    except httpx.ConnectError:
        print(f"\n错误: 无法连接到服务 {BASE_URL}")
        print("请先启动服务: python main.py")
    except Exception as e:
        print(f"\n执行出错: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    asyncio.run(main())
