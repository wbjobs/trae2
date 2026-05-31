import os
import time
import logging
import asyncio
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from typing import Tuple, Optional, Dict, Any, List, Callable
from dataclasses import dataclass, field

import pdfplumber
from docx import Document as DocxDocument
from config import settings

logger = logging.getLogger(__name__)


@dataclass
class ParseProgress:
    current_page: int = 0
    total_pages: int = 0
    current_block: int = 0
    total_blocks: int = 0
    text_length: int = 0
    is_completed: bool = False
    error: Optional[str] = None


class DocumentParser:
    """文档解析模块 - 支持PDF和Word文档解析（优化版）"""

    def __init__(self):
        self.allowed_extensions = settings.ALLOWED_EXTENSIONS
        self.parse_timeout = getattr(settings, 'PARSE_TIMEOUT', 300)
        self.max_pages_per_chunk = getattr(settings, 'MAX_PAGES_PER_CHUNK', 50)
        self.max_text_length = getattr(settings, 'MAX_TEXT_LENGTH', 100000)
        self._executor = ThreadPoolExecutor(max_workers=2)
        logger.info(
            f"文档解析模块初始化完成, 解析超时: {self.parse_timeout}s, "
            f"分页处理大小: {self.max_pages_per_chunk}"
        )

    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """验证文件是否符合要求"""
        ext = os.path.splitext(filename)[1].lower()
        if ext not in self.allowed_extensions:
            return False, f"不支持的文件格式: {ext}, 支持格式: {self.allowed_extensions}"
        if file_size > settings.MAX_UPLOAD_SIZE:
            return False, f"文件大小超过限制: {file_size} > {settings.MAX_UPLOAD_SIZE}"
        return True, None

    def _parse_pdf_with_timeout(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[ParseProgress], None]] = None
    ) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """带超时控制的PDF解析"""
        text_parts: List[str] = []
        metadata: Dict[str, Any] = {}
        page_count = 0
        progress = ParseProgress()

        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                metadata["page_count"] = page_count
                metadata["title"] = pdf.metadata.get("Title", "")
                metadata["author"] = pdf.metadata.get("Author", "")
                metadata["creator"] = pdf.metadata.get("Creator", "")
                metadata["creation_date"] = str(pdf.metadata.get("CreationDate", ""))

                progress.total_pages = page_count
                progress.total_blocks = (page_count + self.max_pages_per_chunk - 1) // self.max_pages_per_chunk

                for chunk_start in range(0, page_count, self.max_pages_per_chunk):
                    chunk_end = min(chunk_start + self.max_pages_per_chunk, page_count)
                    progress.current_block = chunk_start // self.max_pages_per_chunk + 1

                    for page_num in range(chunk_start, chunk_end):
                        progress.current_page = page_num + 1

                        try:
                            page = pdf.pages[page_num]
                            page_text = page.extract_text() or ""

                            if page_text.strip():
                                text_parts.append(f"--- 第{page_num + 1}页 ---\n{page_text}")

                            tables = page.extract_tables()
                            if tables:
                                for table_idx, table in enumerate(tables):
                                    table_text = self._table_to_text(table)
                                    if table_text.strip():
                                        text_parts.append(f"\n[表格 {page_num + 1}-{table_idx}]\n{table_text}")

                        except Exception as page_error:
                            logger.warning(f"PDF第{page_num + 1}页解析失败: {page_error}")
                            continue

                    progress.text_length = sum(len(t) for t in text_parts)
                    if progress_callback:
                        progress_callback(progress)

                    if progress.text_length >= self.max_text_length:
                        logger.info(f"文本长度已达上限: {self.max_text_length}，提前结束解析")
                        break

            full_text = "\n\n".join(text_parts)
            progress.is_completed = True

            if progress_callback:
                progress_callback(progress)

            logger.info(f"PDF解析完成: {file_path}, 页数: {page_count}, 字数: {len(full_text)}")
            return full_text, metadata

        except Exception as e:
            progress.error = str(e)
            if progress_callback:
                progress_callback(progress)
            logger.error(f"PDF解析失败: {file_path}, 错误: {str(e)}")
            return None, {"error": str(e)}

    def parse_pdf(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[ParseProgress], None]] = None
    ) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """解析PDF文档（带超时控制）"""
        loop = asyncio.new_event_loop()
        try:
            future = self._executor.submit(
                self._parse_pdf_with_timeout, file_path, progress_callback
            )
            return future.result(timeout=self.parse_timeout)
        except FuturesTimeoutError:
            error_msg = f"PDF解析超时: {self.parse_timeout}s"
            logger.error(error_msg)
            return None, {"error": error_msg}
        except Exception as e:
            logger.error(f"PDF解析异常: {str(e)}")
            return None, {"error": str(e)}
        finally:
            loop.close()

    def _parse_docx_with_timeout(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[ParseProgress], None]] = None
    ) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """带超时控制的Word文档解析"""
        text_parts: List[str] = []
        metadata: Dict[str, Any] = {}
        progress = ParseProgress()

        try:
            doc = DocxDocument(file_path)

            core_props = doc.core_properties
            metadata["title"] = core_props.title or ""
            metadata["author"] = core_props.author or ""
            metadata["subject"] = core_props.subject or ""
            metadata["created"] = str(core_props.created) if core_props.created else ""
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)

            progress.total_pages = 1
            progress.total_blocks = (len(doc.paragraphs) + 100) // 100

            for para_idx, para in enumerate(doc.paragraphs):
                if para_idx % 100 == 0 and progress_callback:
                    progress.current_block = para_idx // 100 + 1
                    progress.current_page = 1
                    progress.text_length = sum(len(t) for t in text_parts)
                    progress_callback(progress)

                if para.text.strip():
                    style = para.style.name if para.style else ""
                    if style in ["Heading 1", "Heading 2", "Heading 3"]:
                        text_parts.append(f"\n## {para.text}\n")
                    else:
                        text_parts.append(para.text)

            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                table_text = self._table_to_text(table_data)
                if table_text.strip():
                    text_parts.append(f"\n[表格 {table_idx + 1}]\n{table_text}")

            full_text = "\n".join(text_parts)
            progress.is_completed = True
            progress.text_length = len(full_text)

            if progress_callback:
                progress_callback(progress)

            logger.info(f"Word文档解析完成: {file_path}, 段落数: {len(doc.paragraphs)}, 字数: {len(full_text)}")
            return full_text, metadata

        except Exception as e:
            progress.error = str(e)
            if progress_callback:
                progress_callback(progress)
            logger.error(f"Word文档解析失败: {file_path}, 错误: {str(e)}")
            return None, {"error": str(e)}

    def parse_docx(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[ParseProgress], None]] = None
    ) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """解析Word文档（带超时控制）"""
        loop = asyncio.new_event_loop()
        try:
            future = self._executor.submit(
                self._parse_docx_with_timeout, file_path, progress_callback
            )
            return future.result(timeout=self.parse_timeout)
        except FuturesTimeoutError:
            error_msg = f"Word文档解析超时: {self.parse_timeout}s"
            logger.error(error_msg)
            return None, {"error": error_msg}
        except Exception as e:
            logger.error(f"Word文档解析异常: {str(e)}")
            return None, {"error": str(e)}
        finally:
            loop.close()

    def parse_document(
        self,
        file_path: str,
        file_type: str,
        progress_callback: Optional[Callable[[ParseProgress], None]] = None
    ) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """根据文件类型解析文档"""
        file_type = file_type.lower()

        if file_type == "pdf" or file_path.lower().endswith(".pdf"):
            return self.parse_pdf(file_path, progress_callback)
        elif file_type in ["docx", "doc"] or file_path.lower().endswith((".docx", ".doc")):
            return self.parse_docx(file_path, progress_callback)
        else:
            logger.error(f"不支持的文档类型: {file_type}")
            return None, {"error": f"不支持的文档类型: {file_type}"}

    def _table_to_text(self, table: list) -> str:
        """将表格数据转换为文本格式"""
        if not table:
            return ""

        lines = []
        for row in table:
            row_text = " | ".join(str(cell).strip() for cell in row if str(cell).strip())
            if row_text:
                lines.append(row_text)

        return "\n".join(lines)

    def clean_text(self, raw_text: str, max_length: Optional[int] = None) -> str:
        """清洗解析后的文本"""
        if not raw_text:
            return ""

        if max_length is None:
            max_length = self.max_text_length

        lines = raw_text.split("\n")
        cleaned_lines = []
        current_length = 0

        for line in lines:
            line = line.strip()
            if line:
                line = " ".join(line.split())
                cleaned_lines.append(line)
                current_length += len(line)

                if current_length >= max_length:
                    logger.debug(f"文本清洗达到长度限制: {max_length}")
                    break

        cleaned_text = "\n".join(cleaned_lines)
        logger.debug(f"文本清洗完成, 原长度: {len(raw_text)}, 清洗后: {len(cleaned_text)}")
        return cleaned_text

    def extract_text_chunks(
        self,
        text: str,
        chunk_size: int = 4000,
        overlap: int = 200
    ) -> List[str]:
        """将长文本分割成重叠块（用于AI处理）"""
        if not text or len(text) <= chunk_size:
            return [text] if text else []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = min(start + chunk_size, text_length)
            chunk = text[start:end]
            chunks.append(chunk)

            if end >= text_length:
                break

            start = end - overlap

        logger.debug(f"文本分块完成: {len(chunks)}块, 每块最大{chunk_size}字符, 重叠{overlap}字符")
        return chunks


document_parser = DocumentParser()
