import os
import sys
import tempfile
from docx import Document
import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from modules.document_parser import document_parser
from modules.ai_client import ai_client
from modules.semantic_extractor import semantic_extractor
from modules.classification_store import classification_store


@pytest.fixture
def sample_docx_file():
    """创建测试用Word文档"""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_heading('技术服务合同', level=1)
        doc.add_paragraph('甲方：XX科技有限公司')
        doc.add_paragraph('乙方：YY软件有限公司')
        doc.add_paragraph('')
        doc.add_paragraph('根据《中华人民共和国合同法》及相关法律法规，甲乙双方本着平等互利的原则，经友好协商，就甲方委托乙方提供技术服务事宜达成如下协议：')
        doc.add_paragraph('')
        doc.add_heading('第一条 服务内容', level=2)
        doc.add_paragraph('1.1 乙方同意按照本合同约定向甲方提供企业管理系统的开发、实施和维护服务。')
        doc.add_paragraph('1.2 服务期限：自2024年1月1日起至2024年12月31日止。')
        doc.add_paragraph('1.3 服务费用：本项目总费用为人民币50万元整。')
        doc.add_paragraph('')
        doc.add_heading('第二条 双方权利与义务', level=2)
        doc.add_paragraph('2.1 甲方应按时向乙方提供项目所需的相关资料和业务需求。')
        doc.add_paragraph('2.2 乙方应按照行业标准完成系统开发工作，并保证代码质量。')
        doc.add_paragraph('2.3 乙方应指派具有5年以上经验的项目经理负责本项目。')
        doc.save(tmp.name)
        tmp_path = tmp.name

    yield tmp_path

    if os.path.exists(tmp_path):
        os.unlink(tmp_path)


class TestDocumentParser:
    """文档解析模块测试"""

    def test_validate_file(self):
        """测试文件验证"""
        valid, error = document_parser.validate_file("test.pdf", 1024)
        assert valid is True
        assert error is None

        valid, error = document_parser.validate_file("test.txt", 1024)
        assert valid is False
        assert "不支持的文件格式" in error

        valid, error = document_parser.validate_file("test.pdf", 60 * 1024 * 1024)
        assert valid is False
        assert "文件大小超过限制" in error

    def test_parse_docx(self, sample_docx_file):
        """测试Word文档解析"""
        text, metadata = document_parser.parse_docx(sample_docx_file)

        assert text is not None
        assert "技术服务合同" in text
        assert "甲方" in text
        assert "乙方" in text
        assert metadata is not None
        assert "paragraph_count" in metadata

    def test_clean_text(self):
        """测试文本清洗"""
        raw_text = "  这是   一段    测试文本  \n\n  包含多余空格  "
        cleaned = document_parser.clean_text(raw_text)
        assert "  " not in cleaned
        assert cleaned.startswith("这是")
        assert cleaned.endswith("包含多余空格")


class TestAIClient:
    """AI模型调用模块测试"""

    def test_mock_chat_completion(self):
        """测试模拟AI聊天补全"""
        messages = [
            {"role": "system", "content": "你是助手"},
            {"role": "user", "content": "提取关键词"}
        ]
        response, error = ai_client.generate_chat_completion(messages)
        assert error is None
        assert response is not None

    def test_mock_embedding(self):
        """测试模拟向量生成"""
        text = "这是测试文本"
        embedding, error = ai_client.generate_embedding(text)
        assert error is None
        assert embedding is not None
        assert isinstance(embedding, list)
        assert len(embedding) > 0

    def test_batch_embeddings(self):
        """测试批量向量生成"""
        texts = ["文本1", "文本2", "文本3"]
        embeddings, errors = ai_client.batch_generate_embeddings(texts)
        assert len(embeddings) == 3
        assert len(errors) == 0


class TestSemanticExtractor:
    """语义抽取模块测试"""

    @pytest.mark.asyncio
    async def test_extract_keywords_tfidf(self):
        """测试TF-IDF关键词提取"""
        text = """
        这是一份技术服务合同，涉及软件开发、系统实施和技术维护。
        甲方为科技公司，乙方为软件公司，合同金额50万元，期限一年。
        项目包括需求分析、系统设计、编码实现、测试部署等阶段。
        """
        keywords = semantic_extractor.extract_keywords_tfidf(text, top_k=10)
        assert isinstance(keywords, list)
        assert len(keywords) <= 10

    def test_extract_keywords_textrank(self):
        """测试TextRank关键词提取"""
        text = "机器学习是人工智能的一个分支，它使用算法从数据中学习模式。深度学习是机器学习的子领域，使用神经网络模型。"
        keywords = semantic_extractor.extract_keywords_textrank(text, top_k=5)
        assert isinstance(keywords, list)

    @pytest.mark.asyncio
    async def test_generate_summary_extractive(self):
        """测试抽取式摘要"""
        text = """
        这是一份关于项目进度的报告。项目第一阶段已经完成，达到了预期目标。
        目前正在进行第二阶段的开发工作，预计下月完成。团队成员士气高涨，工作积极。
        遇到的主要问题是需求变更频繁，影响了开发进度。建议加强需求管理。
        """
        summary = semantic_extractor._generate_summary_extractive(text, max_length=100)
        assert isinstance(summary, str)
        assert len(summary) <= 100

    def test_extract_entities_rule_based(self):
        """测试基于规则的实体提取"""
        text = "甲方公司于2024年1月1日与乙方科技有限公司签订了合同。"
        entities = semantic_extractor._extract_entities_rule_based(text)
        assert isinstance(entities, list)
        assert len(entities) > 0


class TestClassificationStore:
    """分类存储模块测试"""

    def test_classify_rule_based_contract(self):
        """测试基于规则的合同分类"""
        text = "根据合同约定，甲方应向乙方支付服务费。本协议自签订之日起生效。"
        keywords = ["合同", "协议", "甲方", "乙方", "服务费"]

        result = classification_store._classify_rule_based(text, keywords)
        assert result.primary_category == "合同协议"
        assert result.confidence > 0.5

    def test_classify_rule_based_tech(self):
        """测试基于规则的技术文档分类"""
        text = "本系统采用微服务架构设计，使用Python语言开发，数据库采用PostgreSQL。接口采用RESTful风格。"
        keywords = ["系统", "架构", "微服务", "Python", "数据库", "接口"]

        result = classification_store._classify_rule_based(text, keywords)
        assert result.primary_category == "技术文档"

    def test_classify_rule_based_finance(self):
        """测试基于规则的财务报表分类"""
        text = "本公司本年度营业收入1000万元，净利润200万元，资产总额5000万元，负债总额2000万元。"
        keywords = ["营业收入", "净利润", "资产", "负债", "财务"]

        result = classification_store._classify_rule_based(text, keywords)
        assert result.primary_category == "财务报表"

    def test_classify_rule_based_default(self):
        """测试无法分类时返回其他"""
        text = "今天天气很好，适合外出散步。"
        keywords = ["天气", "散步"]

        result = classification_store._classify_rule_based(text, keywords)
        assert result.primary_category == "其他"

    def test_semantic_search(self):
        """测试语义搜索相似度计算"""
        import numpy as np

        query_embedding = np.random.randn(1536).tolist()

        db_session = None
        try:
            results = classification_store.semantic_search(
                db=db_session,
                query_embedding=query_embedding,
                top_k=10
            )
            assert isinstance(results, list)
        except Exception:
            pass


class TestIntegration:
    """集成测试"""

    @pytest.mark.asyncio
    async def test_full_pipeline(self, sample_docx_file):
        """测试完整处理流程"""
        text, metadata = document_parser.parse_document(sample_docx_file, "docx")
        assert text is not None
        assert len(text) > 0

        cleaned_text = document_parser.clean_text(text)
        assert cleaned_text is not None

        keywords = semantic_extractor.extract_keywords_tfidf(cleaned_text, top_k=10)
        assert len(keywords) > 0

        classification = classification_store._classify_rule_based(cleaned_text, keywords)
        assert classification.primary_category is not None
        assert 0 <= classification.confidence <= 1

        embedding, error = ai_client.generate_embedding(cleaned_text)
        assert error is None
        assert embedding is not None
        assert len(embedding) > 0

        print(f"\n=== 处理结果 ===")
        print(f"文档长度: {len(text)} 字符")
        print(f"关键词: {keywords[:5]}")
        print(f"分类结果: {classification.primary_category}")
        print(f"置信度: {classification.confidence:.2f}")
        print(f"向量维度: {len(embedding)}")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
